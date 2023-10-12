#!/usr/bin/env python

import json
from sys import argv
from docx import Document
from typing import List, Tuple

placeholder_mapping = {
    '[placeholder1]': 'algosec_requestorName',
    '[placeholder2]': 'algosec_requestorID',
    '[placeholder3]': 'algosec_requestorEmail',
    '[placeholder4]': 'algosec_requestorDepartment',
    '[placeholder5]': 'algosec_technicalContactName',
    '[placeholder6]': 'algosec_technicalContactID',
    '[placeholder7]': 'algosec_technicalContactEmail',
    '[placeholder8]': 'algosec_technicalContactDepartment',
}

def get_placeholders(table) -> List[Tuple[str, int]]:
    row = 0
    results = []
    while True:
        try:
            text = table.cell(row, 1).text
            if 'placeholder' in text:
                results.append((text, row))
        except IndexError:
            break
        else:
            row += 1
    return results


def add_traffic_lines(table) -> None:
    """Adds traffic lines to Word document"""
    keys = ['Source', 'Destination', 'ServiceApplication']
    results = []
    for traffic_line in traffic_lines:
        result_dict = {}
        for key in keys:
            if isinstance(traffic_line[key], dict) or isinstance(traffic_line[key], str):
                traffic_line[key] = [traffic_line[key]]
        result_dict['Source'] = '\n'.join(f"{source['Value']} {source['source_name']['Value']}" if 'source_name' in source else source['Value'] for source in traffic_line['Source'])
        result_dict['Destination'] = '\n'.join(f"{dest['Value']} {dest['destination_name']['Value']}" if 'destination_name' in dest else dest['Value'] for dest in traffic_line['Destination'])
        result_dict['ServiceApplication'] = '\n'.join(f"{svc['Value']}/{svc['service_description']['Value']}" if 'service_description' in svc else svc['Value'] for svc in traffic_line['ServiceApplication'])
        result_dict['ShortComment'] = f"{traffic_line['ShortComment']} [{traffic_line['DataConfidentiality']}]" if 'ShortComment' in traffic_line and traffic_line['ShortComment'] else f"[{traffic_line['DataConfidentiality']}]"
        results.append(result_dict)

    for row, line in enumerate(results):
        table.add_row()
        row += 1
        table.cell(row, 0).text = line['Source']
        table.cell(row, 1).text = line['Destination']
        table.cell(row, 2).text = line['ServiceApplication']
        table.cell(row, 3).text = line['ShortComment']

def main() -> None:
    col_index = 1
    for table_index, table in enumerate(document.tables[:-1]):  # skip last table
        for placeholder, row_index in get_placeholders(table):
            try:
                mapped_key = placeholder_mapping[placeholder]
                new_text = algosec_parameters[mapped_key]
            except KeyError:
                new_text = ''
            document.tables[table_index].cell(row_index, col_index).text = new_text

    add_traffic_lines(document.tables[-1])  # update last table from template document
    document.save(file_path.replace('_template', ''))  # Save to FWCR.docx file


if __name__ == "__main__":
    with open(argv[1], 'r') as fp:
        algosec_parameters = json.load(fp)

    traffic_lines = json.loads(algosec_parameters['algosec_traffic_lines'])
    if isinstance(traffic_lines, dict):
        traffic_lines = [traffic_lines]

    file_path = argv[2]
    document = Document(file_path)

    main()
